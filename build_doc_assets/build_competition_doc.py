from pathlib import Path
from zipfile import ZipFile
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(r"D:\STUDY\hrrr-radar-monitor-system")
DOCS = BASE / "docs"
ASSETS = BASE / "build_doc_assets"
ASSETS.mkdir(exist_ok=True)
SRC_DOCX = next(DOCS.glob("比赛*.docx"))
OUT_DOCX = DOCS / "比赛提交文档-多模态非接触式睡眠监测系统.docx"

ACCENT = RGBColor(31, 78, 121)
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(20, 35, 54)
MUTED = RGBColor(90, 105, 120)


def extract_template_images():
    with ZipFile(SRC_DOCX) as zf:
        for name in zf.namelist():
            if name.startswith("word/media/"):
                (ASSETS / Path(name).name).write_bytes(zf.read(name))


def load_font(size, bold=False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def wrap_by_px(draw, text, font, max_width):
    lines = []
    for raw in str(text).split("\n"):
        line = ""
        for ch in raw:
            test = line + ch
            width = draw.textbbox((0, 0), test, font=font)[2]
            if width > max_width and line:
                lines.append(line)
                line = ch
            else:
                line = test
        if line:
            lines.append(line)
    return lines or [""]


def rounded_box(draw, xy, fill, outline, radius=18, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, color="#38658A", width=4):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    import math
    angle = math.atan2(y2 - y1, x2 - x1)
    length = 16
    spread = 0.42
    points = [
        end,
        (x2 - length * math.cos(angle - spread), y2 - length * math.sin(angle - spread)),
        (x2 - length * math.cos(angle + spread), y2 - length * math.sin(angle + spread)),
    ]
    draw.polygon(points, fill=color)


def draw_label(draw, xy, text, font, color="#142336", anchor="mm", max_width=None, line_gap=6):
    x, y = xy
    if max_width:
        lines = wrap_by_px(draw, text, font, max_width)
    else:
        lines = str(text).split("\n")
    heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total = sum(heights) + line_gap * (len(lines) - 1)
    cy = y - total / 2
    for line, h in zip(lines, heights):
        draw.text((x, cy), line, font=font, fill=color, anchor="ma")
        cy += h + line_gap


def create_diagram(path, title, boxes, links=None, size=(1500, 900)):
    im = Image.new("RGB", size, "#FFFFFF")
    draw = ImageDraw.Draw(im)
    title_font = load_font(38, True)
    body_font = load_font(24)
    small_font = load_font(20)
    draw.text((60, 42), title, font=title_font, fill="#1F4E79")
    draw.line((60, 92, size[0] - 60, 92), fill="#B8D3E8", width=3)
    if links:
        for a, b in links:
            ax1, ay1, ax2, ay2 = boxes[a]["xy"]
            bx1, by1, bx2, by2 = boxes[b]["xy"]
            start = (ax2, (ay1 + ay2) // 2)
            end = (bx1, (by1 + by2) // 2)
            if abs(start[1] - end[1]) < 12:
                arrow(draw, start, end)
            else:
                mid = ((start[0] + end[0]) // 2, start[1])
                draw.line([start, mid, (mid[0], end[1])], fill="#38658A", width=4)
                arrow(draw, (mid[0], end[1]), end)
    for key, item in boxes.items():
        xy = item["xy"]
        rounded_box(draw, xy, item.get("fill", "#EAF2F8"), item.get("outline", "#6BAED6"))
        cx = (xy[0] + xy[2]) // 2
        if item.get("desc"):
            height = xy[3] - xy[1]
            title_y = xy[1] + int(height * 0.35)
            desc_y = xy[1] + int(height * 0.68)
            draw_label(draw, (cx, title_y), item["title"], body_font, "#142336", max_width=xy[2] - xy[0] - 32)
            draw_label(draw, (cx, desc_y), item["desc"], small_font, "#4B6175", max_width=xy[2] - xy[0] - 32)
        else:
            cy = (xy[1] + xy[3]) // 2
            draw_label(draw, (cx, cy), item["title"], body_font, "#142336", max_width=xy[2] - xy[0] - 32)
    im.save(path, quality=95)


def make_custom_figures():
    create_diagram(
        ASSETS / "diagram_system_architecture.png",
        "系统总体架构：边缘感知、真实后端、Web 与 Android 告警闭环",
        {
            "sensors": {"xy": (80, 180, 360, 360), "title": "多模态传感", "desc": "雷达 / 麦克风 / IMU / AHT20", "fill": "#F0F7FF"},
            "m55": {"xy": (460, 160, 760, 380), "title": "Edgi Talk M55", "desc": "RT-Thread\n小智语音 + 呼噜推理 + SOS UI", "fill": "#EAF2F8"},
            "radar": {"xy": (460, 470, 760, 650), "title": "毫米波雷达板", "desc": "心率、呼吸率、距离、在床状态", "fill": "#EEF9F0"},
            "backend": {"xy": (870, 250, 1150, 520), "title": "FastAPI 真实后端", "desc": "状态聚合 / 事件融合\n疑似呼吸暂停判断", "fill": "#FFF6E5"},
            "web": {"xy": (1250, 160, 1450, 330), "title": "Web 看护端", "desc": "驾驶舱\n预警中心\n环境分析", "fill": "#F3ECFF"},
            "app": {"xy": (1250, 460, 1450, 630), "title": "Android App", "desc": "后台轮询\n本地通知\n已处理", "fill": "#FDEFF4"},
        },
        [("sensors", "m55"), ("sensors", "radar"), ("m55", "backend"), ("radar", "backend"), ("backend", "web"), ("backend", "app")],
    )
    create_diagram(
        ASSETS / "diagram_audio_hub.png",
        "共享麦克风采集中心：避免语音与呼噜检测抢占 mic0",
        {
            "mic": {"xy": (80, 300, 310, 480), "title": "mic0", "desc": "16 kHz 双声道 PCM", "fill": "#EEF9F0"},
            "hub": {"xy": (430, 230, 740, 550), "title": "音频采集中心", "desc": "唯一打开麦克风\n非阻塞环形缓冲\n消费者独立丢帧", "fill": "#EAF2F8"},
            "wake": {"xy": (900, 140, 1250, 290), "title": "唤醒词消费者", "desc": "右声道单声道\n待机常驻", "fill": "#F0F7FF"},
            "voice": {"xy": (900, 360, 1250, 510), "title": "小智语音上传", "desc": "按钮/唤醒后开启\nSTT + TTS", "fill": "#FFF6E5"},
            "snore": {"xy": (900, 580, 1250, 730), "title": "呼噜检测消费者", "desc": "2 秒窗口\n1 秒推理", "fill": "#FDEFF4"},
        },
        [("mic", "hub"), ("hub", "wake"), ("hub", "voice"), ("hub", "snore")],
    )
    create_diagram(
        ASSETS / "diagram_emergency_flow.png",
        "紧急求助闭环：语音/IMU 触发，本地与远程同步处理",
        {
            "trigger": {"xy": (80, 240, 310, 430), "title": "触发源", "desc": "救命 / 需要帮助\nIMU 摇晃 / 打翻", "fill": "#FDEFF4"},
            "board": {"xy": (420, 210, 690, 460), "title": "开发板 SOS", "desc": "红色界面\n报警声\n解除按钮", "fill": "#FFE8E8"},
            "api": {"xy": (800, 240, 1060, 430), "title": "后端事件", "desc": "/emergency\ncritical\nactive_emergency", "fill": "#FFF6E5"},
            "web": {"xy": (1160, 120, 1420, 300), "title": "Web 预警中心", "desc": "置顶告警\n处理动作", "fill": "#F3ECFF"},
            "app": {"xy": (1160, 430, 1420, 610), "title": "安卓通知", "desc": "5 秒轮询\n锁屏提醒", "fill": "#EAF2F8"},
            "resolve": {"xy": (520, 620, 960, 770), "title": "人工确认后解除", "desc": "开发板按钮 / 前端处理 / App 已处理 -> /emergency/resolve", "fill": "#EEF9F0"},
        },
        [("trigger", "board"), ("board", "api"), ("api", "web"), ("api", "app"), ("web", "resolve"), ("app", "resolve"), ("resolve", "api")],
    )
    create_diagram(
        ASSETS / "diagram_apnea_fusion.png",
        "疑似呼吸暂停融合：雷达低呼吸 + 呼噜/恢复性声音证据",
        {
            "radar": {"xy": (80, 190, 360, 390), "title": "雷达窗口", "desc": "目标在床\n呼吸率连续低/丢失", "fill": "#EAF2F8"},
            "snore": {"xy": (80, 500, 360, 700), "title": "声音窗口", "desc": "呼噜增强\ndBFS 变化\n恢复性声音", "fill": "#FDEFF4"},
            "gate": {"xy": (520, 270, 850, 620), "title": "规则融合器", "desc": "排除：无人、离线、未静止\n时间窗：30-60 秒\n去重：同一段只记一次", "fill": "#FFF6E5"},
            "event": {"xy": (1030, 350, 1370, 540), "title": "suspected_apnea", "desc": "warning / critical\n事件流 + App 通知\n非医疗诊断", "fill": "#FFE8E8"},
        },
        [("radar", "gate"), ("snore", "gate"), ("gate", "event")],
    )
    create_diagram(
        ASSETS / "diagram_backend_interfaces.png",
        "真实后端接口与数据对象",
        {
            "hw": {"xy": (70, 190, 440, 430), "title": "真实硬件上报", "desc": "snore heartbeat\nenvironment heartbeat\nemergency event", "fill": "#EAF2F8"},
            "udp": {"xy": (70, 520, 440, 710), "title": "雷达 UDP", "desc": "raw frame\nheart / breath / distance", "fill": "#EEF9F0"},
            "state": {"xy": (560, 250, 930, 650), "title": "运行状态层", "desc": "timeline\nsleep_events\ndevices\nenvironment\nactive emergency", "fill": "#FFF6E5"},
            "ui": {"xy": (1080, 190, 1430, 430), "title": "Web / App 读取", "desc": "status\ntimeline\nsleep overview", "fill": "#F3ECFF"},
            "resolve": {"xy": (1080, 520, 1430, 710), "title": "人工处理", "desc": "emergency resolve\n处理人 / 备注", "fill": "#FDEFF4"},
        },
        [("hw", "state"), ("udp", "state"), ("state", "ui"), ("resolve", "state")],
    )
    create_diagram(
        ASSETS / "diagram_board_m55_architecture.png",
        "Edgi Talk M55 小智/呼噜/SOS 板系统架构",
        {
            "mic": {"xy": (70, 170, 320, 340), "title": "PDM 麦克风", "desc": "语音 / 唤醒词 / 呼噜\n共享音频输入", "fill": "#EEF9F0"},
            "imu": {"xy": (70, 410, 320, 570), "title": "IMU", "desc": "摇晃 / 打翻\n无法呼救场景触发", "fill": "#FDEFF4"},
            "env": {"xy": (70, 640, 320, 790), "title": "AHT20", "desc": "温度 / 湿度\n环境评分输入", "fill": "#FFF6E5"},
            "hub": {"xy": (450, 200, 750, 390), "title": "音频采集中心", "desc": "唯一打开 mic0\n非阻塞环形缓冲\n消费者独立丢帧", "fill": "#EAF2F8"},
            "m55": {"xy": (450, 500, 750, 720), "title": "PSOC Edge M55", "desc": "RT-Thread\n小智语音\nint8 呼噜推理\nSOS 状态机", "fill": "#F0F7FF"},
            "ui": {"xy": (900, 210, 1180, 400), "title": "屏幕与按键", "desc": "主页 / 呼噜状态\nSOS 解除按钮\n闹钟界面", "fill": "#F3ECFF"},
            "wifi": {"xy": (900, 520, 1180, 710), "title": "AIROC Wi-Fi", "desc": "HTTP 上报\n/hardware/*\n/emergency", "fill": "#EEF9F0"},
            "backend": {"xy": (1280, 360, 1460, 570), "title": "真实后端", "desc": "事件流\n设备在线\n环境趋势", "fill": "#FFF6E5"},
        },
        [("mic", "hub"), ("hub", "m55"), ("imu", "m55"), ("env", "m55"), ("m55", "ui"), ("m55", "wifi"), ("wifi", "backend")],
    )
    create_diagram(
        ASSETS / "diagram_board_radar_architecture.png",
        "毫米波雷达生命体征板系统架构",
        {
            "radar": {"xy": (70, 230, 360, 470), "title": "BGT60TR13C 雷达", "desc": "毫米波 FMCW\n胸腔微动 / 距离信息\n非接触采集", "fill": "#EAF2F8"},
            "mcu": {"xy": (510, 190, 830, 510), "title": "PSoC 6 / 雷达任务", "desc": "雷达初始化\n帧读取与缓存\n状态机控制\nLED 指示", "fill": "#F0F7FF"},
            "wifi": {"xy": (980, 190, 1280, 510), "title": "AIROC Wi-Fi", "desc": "UDP 数据链路\n接收 enable 命令\n发送雷达帧", "fill": "#EEF9F0"},
            "backend": {"xy": (1040, 610, 1380, 770), "title": "后端雷达处理", "desc": "距离 FFT\n相位提取\n心率 / 呼吸率 / 在床判断", "fill": "#FFF6E5"},
            "state": {"xy": (510, 620, 830, 790), "title": "板载状态", "desc": "发送中 / 暂停\nWi-Fi 连接\n雷达初始化结果", "fill": "#FDEFF4"},
        },
        [("radar", "mcu"), ("mcu", "wifi"), ("wifi", "backend"), ("mcu", "state")],
    )
    create_diagram(
        ASSETS / "diagram_board_env_architecture.png",
        "M33 / AHT20 环境采集板系统架构",
        {
            "aht": {"xy": (80, 270, 350, 500), "title": "AHT20 温湿度传感器", "desc": "I2C 总线\n温度 / 湿度\nsensor_ok 状态", "fill": "#FFF6E5"},
            "m33": {"xy": (520, 210, 850, 560), "title": "Cortex-M33 环境任务", "desc": "周期采样\n异常检测\n共享数据缓存\n降低串口噪声", "fill": "#EAF2F8"},
            "m55": {"xy": (1010, 210, 1320, 560), "title": "M55 上报任务", "desc": "读取最新样本\n组装 JSON\nHTTP 心跳", "fill": "#F0F7FF"},
            "backend": {"xy": (1010, 650, 1320, 790), "title": "环境分析", "desc": "温湿度趋势\n分贝\n睡眠环境评分", "fill": "#EEF9F0"},
        },
        [("aht", "m33"), ("m33", "m55"), ("m55", "backend")],
    )


def font_run(run, size=10.5, color=DARK, bold=None):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=8.8, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) <= 14 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(str(text))
    font_run(run, size=size, color=color or DARK, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D9E2EC"):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            cell = row.cells[idx]
            cell.width = Inches(width)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(int(width * 1440)))
            tcW.set(qn("w:type"), "dxa")


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        font_run(run, size=16 if level == 1 else (13 if level == 2 else 11.8), color=BLUE if level <= 2 else ACCENT, bold=True)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(4)
    set_keep_with_next(p)
    return p


def add_para(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21) if indent else Pt(0)
    p.paragraph_format.space_after = Pt(4.5)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    font_run(run, size=10.35)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.line_spacing = 1.08
        run = p.add_run(item)
        font_run(run, size=10)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    font_run(run, size=8.8, color=MUTED)
    return p


def add_figure(doc, filename, caption, width=5.7):
    path = ASSETS / filename
    if not path.exists():
        add_caption(doc, caption + "（图片待补充）")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=8.8, color=RGBColor(255, 255, 255))
        set_cell_shading(table.rows[0].cells[i], "1F4E79")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, size=8.55)
    if widths:
        set_table_widths(table, widths)
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EAF2F8")
    set_table_borders(table, color="B8D3E8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    font_run(r, size=10.3, color=ACCENT, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    rr = p2.add_run(body)
    font_run(rr, size=9.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def page_break(doc):
    doc.add_page_break()


def build_doc():
    extract_template_images()
    make_custom_figures()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.3)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "微软雅黑"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("多模态非接触式睡眠监测系统 · 好好睡觉")
    font_run(footer_run, size=8, color=MUTED)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run("多模态非接触式睡眠监测系统")
    font_run(r, size=25, color=ACCENT, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("基于 PSOC Edge / Edgi Talk 的夜间安全与睡眠看护平台")
    font_run(r, size=13, color=MUTED)
    add_figure(doc, "diagram_system_architecture.png", "封面图 系统整体技术路线", width=5.9)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(4)
    r = p3.add_run("队名：好好睡觉\n成员：陈俊帆、王志华\n赛道：2026 嵌入式系统设计大赛 · 英飞凌赛道")
    font_run(r, size=11.5)
    page_break(doc)

    add_heading(doc, "摘要", 1)
    add_para(doc, "随着居家养老、慢病管理与睡眠健康关注度持续提升，夜间睡眠阶段的呼吸异常、摇晃风险、环境不舒适和无法主动求助等问题逐渐成为家庭看护的薄弱环节。传统摄像头看护存在隐私压力，可穿戴设备又容易因佩戴不适、忘记充电或摘除而中断监测。针对上述问题，本作品提出并实现了一套多模态非接触式睡眠监测系统，面向独居老人、夜间照护、养老机构巡护和睡眠健康辅助观察场景，提供生命体征监测、呼噜识别、疑似呼吸暂停提示、睡眠环境评估和紧急求助告警能力。")
    add_para(doc, "系统以 PSOC Edge / Edgi Talk 开发板为核心边缘终端，运行 RT-Thread，利用麦克风、IMU、AHT20 温湿度传感器和毫米波雷达构成多源感知。边缘侧完成小智语音交互、呼噜 int8 模型推理、摇晃触发、环境上报和本地 SOS 显示；后端通过 FastAPI 汇聚真实开发板数据，结合雷达呼吸率和呼噜强度进行疑似睡眠呼吸暂停融合判断；前端提供睡眠看护驾驶舱、看护预警中心和睡眠环境分析页面；Android 监护人 App 通过前台服务轮询后端，在锁屏或后台状态下接收紧急通知。")
    add_para(doc, "作品重点体现英飞凌赛道关注的多传感器融合、边缘 AI、无线互联、工程样机完整度和展示实用性。系统不作为医疗诊断设备，而作为课程设计与看护辅助展示平台，为家庭和社区照护提供低侵入、可视化、可扩展的夜间安全监测方案。")
    add_callout(doc, "关键词", "PSOC Edge；Edgi Talk；RT-Thread；毫米波雷达；呼噜检测；睡眠呼吸暂停；小智语音；边缘 AI；Android 告警")

    add_heading(doc, "第一部分  赛道适配与作品概述", 1)
    add_heading(doc, "1.1 英飞凌赛道要求响应", 2)
    add_para(doc, "2026 年英飞凌赛道鼓励参赛作品使用 PSOC Edge、PSoC 6 或 XMC7000 等平台，重点考察 MCU/NPU/DSP 结合、多传感器外设、无线连接、工程样机集成度、低功耗与实际落地价值。本作品选择 PSOC Edge / Edgi Talk 作为边缘智能终端，把音频、显示、Wi-Fi、IMU、温湿度和模型推理融合到一个夜间看护样机中，契合“嵌入式物联网”和“边缘人工智能”方向。")
    add_table(doc, ["赛道评价点", "本作品对应设计", "体现方式"], [
        ["PSOC Edge 平台", "Edgi Talk M55 + RT-Thread", "语音、显示、音频采集、边缘推理、Wi-Fi 上报"],
        ["NPU/DSP/边缘 AI", "呼噜 int8 模型边缘推理", "本地完成声音事件识别，减少持续上传原始音频"],
        ["多传感器融合", "雷达、麦克风、IMU、AHT20", "生命体征、声音、动作、环境四类数据互补"],
        ["无线连接", "AIROC Wi-Fi + Android App 远程监听", "真实后端接口、前端驾驶舱、手机通知"],
        ["工程样机完整度", "开发板 UI、报警声、解除按钮、Web 与 App", "形成可演示、可测试、可复现闭环"],
        ["实用性", "夜间安全与睡眠监测", "聚焦居家养老和睡眠健康辅助观察"],
    ], widths=[1.45, 2.25, 2.7])
    add_figure(doc, "infineon_psoc_edge-08.png", "图1 PSOC Edge E84 平台能力概览（来源：英飞凌赛道资料）", width=5.7)
    add_figure(doc, "infineon_requirements-13.png", "图2 英飞凌赛道嵌入式物联网方向评价要求（来源：英飞凌赛道资料）", width=5.7)

    add_heading(doc, "1.2 功能与特性", 2)
    add_bullets(doc, [
        "非接触生命体征监测：毫米波雷达采集人体胸腔微动，后端估计心率、呼吸率、目标距离和在床状态。",
        "呼噜与疑似呼吸暂停提示：M55 共享麦克风采集声音，边缘侧模型实时推理；后端融合雷达低呼吸与呼噜变化生成 suspected_apnea 事件。",
        "语音紧急求助：用户说“救命”“需要帮助”“喘不过气”等关键词时，小智语音触发 SOS、本地报警和后端 critical 事件。",
        "IMU 摇晃/打翻求助：当板子被打翻或检测到摇晃动作时，系统模拟用户无法呼救场景并进入紧急状态。",
        "睡眠环境分析：AHT20 温湿度和麦克风分贝数据形成环境评分，前端给出简洁建议。",
        "多端看护闭环：Web 驾驶舱、预警中心和 Android 监护人 App 同步展示并处理告警。",
    ])
    add_heading(doc, "1.3 应用领域", 2)
    add_para(doc, "本系统主要面向居家养老、社区日间照料、养老院夜间巡护和睡眠健康辅助观察等场景。与摄像头监控相比，毫米波雷达和音频特征监测在隐私保护方面更友好；与单一可穿戴设备相比，非接触式方案降低了使用门槛，适合长期、低打扰运行。系统可以放置在床旁或桌面，用于夜间连续观察，监护人通过 Web 页面或 Android App 接收告警。")
    add_heading(doc, "1.4 主要创新点", 2)
    add_bullets(doc, [
        "将雷达呼吸率与呼噜声证据融合，避免单点雷达异常直接误判为呼吸暂停。",
        "把“小智语音”从普通交互扩展为夜间紧急求助入口，并增加本地 SOS 屏幕与解除按钮。",
        "共享麦克风架构保证呼噜守护和小智语音可同时运行，提升夜间连续监测能力。",
        "开发板端、后端、Web 前端、Android App 形成完整告警闭环，接近真实看护样机。",
        "环境评分将温湿度和环境声纳入睡眠建议，展示从生命体征到环境干预的整体方案。",
    ])

    page_break(doc)
    add_heading(doc, "第二部分  系统组成及硬件设计", 1)
    add_heading(doc, "2.1 整体系统架构", 2)
    add_para(doc, "系统由边缘感知终端、真实数据后端、Web 看护端和 Android 监护人端四层组成。Edgi Talk M55 负责语音、呼噜、IMU、温湿度和屏幕交互；雷达板负责非接触式心率、呼吸率和距离感知；后端统一维护在线状态、事件流、环境数据和疑似呼吸暂停融合结果；前端和 App 面向监护人提供告警展示与处理入口。")
    add_figure(doc, "diagram_system_architecture.png", "图3 本作品系统总体架构与数据闭环", width=6.1)
    add_table(doc, ["数据流", "采集端", "后端处理", "用户侧呈现"], [
        ["生命体征流", "毫米波雷达板 UDP 帧", "距离 FFT、相位提取、心率/呼吸率估计", "驾驶舱生命体征卡片、趋势图"],
        ["声音事件流", "M55 共享麦克风与呼噜模型", "snore_score、dBFS、呼噜扰动事件", "呼噜强度曲线、预警中心动作建议"],
        ["紧急事件流", "小智语音关键词 / IMU 摇晃", "emergency_voice critical 事件与 active 状态", "开发板 SOS、Web 置顶告警、Android 通知"],
        ["环境感知流", "AHT20 温湿度与环境声", "舒适度状态、环境评分与建议", "环境分析页温湿度折线图"],
    ], widths=[1.15, 1.65, 2.0, 1.85])
    add_para(doc, "上述四条数据流并非孤立展示，而是在后端统一时间轴中汇合。雷达负责“人是否在床、呼吸是否稳定”，麦克风负责“是否出现呼噜或求助语音”，IMU 负责“是否存在设备被打翻/摇晃的紧急触发”，环境传感器负责“睡眠环境是否适宜”。多源数据结合后，系统能把普通异常、强告警和设备状态区分开，避免单一传感器误差导致前端刷屏。")
    add_figure(doc, "image2.png", "图4 原理图级硬件连接框图（复用模板原理图并按本项目语境说明）", width=6.1)
    add_heading(doc, "2.2 硬件模块组成", 2)
    add_table(doc, ["模块", "硬件/资源", "接口/协议", "在本项目中的作用"], [
        ["主控与显示", "Edgi Talk M55 / PSOC Edge", "RT-Thread、LVGL、Wi-Fi", "运行小智 UI、SOS 页面、呼噜推理和网络上报"],
        ["音频输入", "PDM 麦克风", "PDM/PCM", "小智语音、唤醒词、求助关键词、呼噜检测共享输入"],
        ["惯性检测", "IMU", "I2C", "检测开发板摇晃/打翻，触发无法呼救场景下的 SOS"],
        ["环境检测", "AHT20/AHT10 兼容温湿度", "I2C + 共享内存/上报任务", "提供睡眠环境评分和温湿度趋势"],
        ["生命体征", "毫米波雷达模块", "UDP 数据链路", "非接触采集心率、呼吸率、距离和在床状态"],
        ["无线通信", "AIROC Wi-Fi", "HTTP / UDP", "开发板和后端之间进行真实数据上报"],
    ], widths=[1.1, 1.7, 1.45, 2.2])
    add_heading(doc, "2.3 关键硬件原理图", 2)
    add_para(doc, "文档复用了原提交材料中的硬件原理图资源，并筛选与当前系统一致的雷达、麦克风和无线通信模块。与当前睡眠看护功能无关的旧控制类、视觉识别类图纸不再放入正文。")
    add_figure(doc, "cy8ckit_radar_crop.png", "图5 毫米波雷达传感器接口原理图，用于非接触式生命体征采集", width=5.6)
    add_figure(doc, "image5.png", "图6 PDM 麦克风输入电路，用于小智语音与呼噜检测", width=5.3)
    add_figure(doc, "cy8ckit_wifi_schematic_crop.png", "图7 AIROC Wi-Fi 模块供电与接口原理图，用于开发板无线通信", width=5.2)
    add_figure(doc, "image1.png", "图8 硬件层、云平台与应用层的系统关系示意", width=5.7)

    add_heading(doc, "2.4 开发板系统架构与板级原理图", 2)
    add_para(doc, "为便于评审理解硬件分工，本节按真实样机中承担不同职责的开发板进行拆分说明。M55 小智板负责夜间人机交互、呼噜边缘推理、SOS 本地处理和温湿度上报；毫米波雷达板负责非接触式生命体征原始数据采集；M33/AHT20 环境采集板负责稳定提供温湿度样本。三块板子通过 Wi-Fi、UDP、HTTP 和后端状态层构成统一系统，但每块板子的采样、推理、通信和显示职责保持相对独立，便于调试和扩展。")
    add_table(doc, ["开发板", "主要硬件资源", "本地固件职责", "向后端输出"], [
        ["Edgi Talk M55 小智/呼噜/SOS 板", "PSOC Edge M55、PDM 麦克风、触摸屏、IMU、AIROC Wi-Fi、蜂鸣/音频播放资源", "共享麦克风采集、小智语音、呼噜 int8 推理、语音/IMU SOS、闹钟、SOS 解除界面", "/hardware/snore-*、/hardware/environment-heartbeat、/emergency、/emergency/resolve"],
        ["毫米波雷达生命体征板", "PSoC 6、BGT60TR13C 雷达、AIROC Wi-Fi、状态 LED", "雷达初始化、UDP 命令监听、雷达帧发送、发送/暂停状态指示", "UDP 雷达帧，后端解析为心率、呼吸率、距离和在床状态"],
        ["M33/AHT20 环境采集板", "Cortex-M33、AHT20/AHT10 温湿度传感器、I2C 接口", "周期读取温湿度、判断传感器是否在线、减少无意义串口打印", "温度、湿度、sensor_ok、最近更新时间"],
    ], widths=[1.55, 1.85, 2.2, 1.55])

    add_heading(doc, "2.4.1 Edgi Talk M55 小智/呼噜/SOS 板", 3)
    add_para(doc, "M55 板是整套作品的边缘智能入口，也是用户最直接接触到的设备。它的核心任务不是单纯采集音频，而是把语音交互、呼噜检测、IMU 异常触发、SOS 页面、闹钟和真实后端上报统一到一个可操作的床旁终端中。为了避免语音服务和呼噜守护抢占麦克风，板端采用“音频采集中心 + 多消费者环形缓冲”的结构，保证呼噜守护常驻时仍然可以通过唤醒词或用户按键进入小智聆听，并继续识别“救命、需要帮助、喘不过气”等紧急话术。")
    add_figure(doc, "diagram_board_m55_architecture.png", "图8-1 Edgi Talk M55 小智/呼噜/SOS 板系统架构", width=6.1)
    add_para(doc, "从信号路径看，PDM 麦克风的 PCM 数据首先进入共享采集中心，再分别送给唤醒词、小智语音上传和呼噜检测模型；IMU 摇晃触发不依赖语音，直接进入 SOS 状态机；AHT20 温湿度样本通过环境任务封装为心跳数据。M55 根据这些输入刷新屏幕状态，在紧急事件中播放报警声并显示“解除紧急状态”按钮，同时通过 Wi-Fi 向后端创建或解除事件。")
    add_figure(doc, "edge_core_pdm_crop.png", "图8-2 Edgi Talk PDM 麦克风相关原理图，用于小智语音和呼噜检测输入（来源：docs/psoc_edge_corev0_2.pdf）", width=5.8)
    add_figure(doc, "edge_basic_imu_aht20_tight.png", "图8-3 Edgi Talk AHT20 与 IMU 相关原理图，支撑环境监测和设备摇晃触发（来源：docs/psoc_edge_basic_V0.2.pdf）", width=5.8)
    add_para(doc, "M55 板的设计重点在于“床旁可用性”。普通状态下屏幕只保留关键状态和控制按钮；呼噜检测命中时显示呼噜相关图像和推理结果；SOS 状态下界面切换为高对比度红色告警，报警声持续更久，并提供本地手动解除入口。该板还负责处理 Wi-Fi 配置、后端 IP 持久化、Flash 写入校验和按键防抖，保证真实演示时不会因为连续点击或网络抖动造成界面锁死。")

    page_break(doc)
    add_heading(doc, "2.4.2 毫米波雷达生命体征板", 3)
    add_para(doc, "雷达板承担非接触式生命体征采集。与 M55 板不同，雷达板不直接做复杂 UI，而是专注于可靠地初始化 BGT60TR13C 雷达、接收后端 enable 命令、发送雷达原始帧，并用 LED 表示发送和暂停状态。后端收到 UDP 数据后再进行距离 FFT、目标距离箱选择、相位提取、呼吸率估计和心率估计。这样的分工可以降低雷达固件复杂度，也便于后端调参和算法迭代。")
    add_table(doc, ["雷达板环节", "输入", "处理", "输出/观察方式"], [
        ["UDP 命令监听", "后端发送 enable / pause 类控制命令", "更新传输状态机，避免上电后无控制就持续发帧", "串口初始化日志与 LED 状态"],
        ["雷达帧采集", "BGT60TR13C 原始帧", "完成雷达设备初始化、帧缓存和基础异常判断", "UDP 原始帧发送到后端"],
        ["状态指示", "Wi-Fi 连接、传输启停、雷达初始化结果", "将高频串口打印改为 LED 指示，降低调试噪声", "发送中/暂停/异常可由 LED 直接观察"],
        ["后端解析", "UDP 雷达帧", "距离 FFT、目标距离箱选择、相位提取和频谱估计", "心率、呼吸率、距离、在床状态"],
    ], widths=[1.2, 1.65, 2.3, 1.35])
    add_figure(doc, "diagram_board_radar_architecture.png", "图8-4 毫米波雷达生命体征板系统架构", width=6.1)
    add_para(doc, "雷达板与后端之间保留 UDP 链路，原因是原始雷达帧吞吐量高，不适合通过低速状态接口传输。蓝牙功能曾用于近场状态实验，但最终监护人报警 App 改为轮询后端，从而摆脱蓝牙距离限制。雷达板当前在作品中主要提供生命体征和在床状态输入，疑似呼吸暂停判断并不直接在雷达板上完成，而是在后端结合呼噜证据进行融合。")
    add_figure(doc, "cy8ckit_block_crop.png", "图8-5 PSoC 6 AI Evaluation Kit 功能框图，展示雷达、Wi-Fi、麦克风和传感器资源（来源：docs/infineon-cy8ckit-062s2-user-guide-usermanual-en.pdf）", width=5.9)
    add_figure(doc, "cy8ckit_radar_crop.png", "图8-6 毫米波雷达传感器接口原理图，支撑雷达帧采集链路（来源：docs/infineon-cy8ckit-062s2-user-guide-usermanual-en.pdf）", width=5.9)
    add_figure(doc, "cy8ckit_wifi_schematic_crop.png", "图8-7 AIROC Wi-Fi 模块供电与接口原理图，支撑 UDP 雷达数据传输（来源：docs/infineon-cy8ckit-062s2-user-guide-usermanual-en.pdf）", width=5.2)
    add_para(doc, "雷达板在系统闭环中的价值主要体现在三个方面：第一，它提供非接触式呼吸和心率趋势，使系统不依赖用户佩戴设备；第二，它提供目标距离和在床状态，帮助后端排除无人场景下的误报警；第三，它为呼噜融合算法提供“呼吸是否持续减弱”的核心证据。由于课程样机场景中雷达姿态和摆放位置会显著影响结果，后端算法会对异常点进行窗口化处理，而不是把单点低呼吸率直接判定为呼吸暂停。")

    add_heading(doc, "2.4.3 M33/AHT20 环境采集板", 3)
    add_para(doc, "环境采集板的职责相对单一，但对睡眠看护展示很重要。温度和湿度不直接触发 SOS，却影响睡眠舒适度和前端环境评分。系统将 AHT20 采样放在轻量周期任务中处理，并把有效样本交给 M55 或后端上报任务；若传感器初始化失败或长时间无有效样本，后端会将环境传感器标记为离线，前端环境分析页也会给出相应提示。")
    add_table(doc, ["字段/状态", "来源", "后端用途", "前端展示"], [
        ["temperature_c", "AHT20 温度采样", "参与环境评分，判断是否偏热或偏冷", "环境分析页当前温度与折线图"],
        ["humidity_rh", "AHT20 湿度采样", "参与环境评分，判断是否偏干或偏湿", "环境分析页当前湿度与折线图"],
        ["sensor_ok", "初始化和周期读数结果", "决定环境传感器在线/离线状态", "设备状态卡片与异常提示"],
        ["last_seen", "后端收到心跳的时间", "避免短时间抖动立刻判定掉线", "刷新后仍能恢复最近状态"],
    ], widths=[1.15, 1.6, 2.15, 1.5])
    add_figure(doc, "diagram_board_env_architecture.png", "图8-8 M33/AHT20 环境采集板系统架构", width=6.1)
    add_figure(doc, "cy8ckit_i2c_crop.png", "图8-9 I2C 接口连接器原理图，可用于 AHT20 温湿度传感器扩展（来源：docs/infineon-cy8ckit-062s2-user-guide-usermanual-en.pdf）", width=5.6)
    add_figure(doc, "cy8ckit_imu_crop.png", "图8-10 六轴 IMU 原理图，展示 I2C 地址与中断连接方式（来源：docs/infineon-cy8ckit-062s2-user-guide-usermanual-en.pdf）", width=5.6)
    add_para(doc, "在真实联调中，环境板曾出现串口周期打印过多的问题，影响观察关键报警日志。因此当前固件保留必要的初始化和异常日志，去掉正常温湿度采样的高频打印。后端仍按心跳时间判断设备在线，前端则在环境分析页展示温湿度折线图、环境分贝、睡眠环境评分和简短建议。")

    page_break(doc)
    add_heading(doc, "第三部分  嵌入式固件设计", 1)
    add_heading(doc, "3.1 RT-Thread 多任务架构", 2)
    add_para(doc, "M55 固件基于 RT-Thread 构建，采用多任务协同方式组织音频采集、语音交互、呼噜检测、环境上报、IMU 监测和 UI 刷新。固件启动后初始化触摸、屏幕、音频、Wi-Fi、环境监测和小智 UI，随后根据 Wi-Fi 与后端状态持续上报真实数据。")
    add_table(doc, ["任务/模块", "主要职责", "关键输出", "异常处理"], [
        ["音频采集中心", "统一打开麦克风并分发 PCM 数据", "语音、唤醒词、呼噜消费者队列", "消费者慢时只丢弃自身旧数据"],
        ["小智语音任务", "WebSocket 语音交互与 STT 文本处理", "普通对话、求助关键词 SOS", "WebSocket busy 时跳过部分音频包"],
        ["呼噜检测任务", "2 秒窗口 int8 模型推理", "snore_score、dbfs、snore_detected", "本地播放期间抑制误报"],
        ["环境上报任务", "读取 AHT20 共享温湿度", "/hardware/environment-heartbeat", "传感器异常时上报 sensor_ok=false"],
        ["IMU 摇晃任务", "检测低重力/撞击/打翻动作", "设备摇晃 SOS", "触发后冷却，避免连续误报"],
        ["UI 与报警任务", "主页、呼噜状态、SOS、闹钟界面", "本地报警和手动解除按钮", "解除后回主页并上报 resolve"],
    ], widths=[1.25, 2.05, 1.75, 1.7])
    add_heading(doc, "3.2 共享麦克风采集中心", 2)
    add_para(doc, "早期实现中，小智语音和呼噜守护分别尝试打开 mic0，容易产生设备抢占、线程退出断言或聆听中无响应的问题。当前方案将麦克风封装为唯一采集中心，持续读取 16 kHz PCM，并使用非阻塞环形缓冲分发给唤醒词、小智语音上传和呼噜检测。这样即使用户开启呼噜守护，小智仍可识别“救命”等紧急关键词。")
    add_figure(doc, "diagram_audio_hub.png", "图9 共享麦克风采集中心数据分发流程", width=6.1)
    add_heading(doc, "3.3 紧急状态与 UI 设计", 2)
    add_para(doc, "开发板本地 UI 不仅展示监测状态，还承担紧急状态的人机交互。语音或 IMU 触发后，屏幕进入 SOS 页面，报警声持续播放，页面显示触发原因和“解除紧急状态”按钮。用户或看护人员确认安全后可在开发板、Web 前端或 Android App 任一端处理告警，后端统一记录解除状态。")
    add_figure(doc, "diagram_emergency_flow.png", "图10 SOS 紧急求助状态机与多端解除闭环", width=6.1)
    add_heading(doc, "3.4 固件存储与可靠性处理", 2)
    add_para(doc, "系统运行中还针对 Flash/Wi-Fi 保存、按钮连击锁死、报警播放误触发呼噜等问题进行了处理。Wi-Fi 配置保存前会进行写入校验；呼噜启停使用生命周期锁和线程退出等待；本地 TTS、闹钟或 SOS 播放期间暂时抑制呼噜告警，避免设备自身声音被模型识别为呼噜。")
    add_table(doc, ["可靠性问题", "工程处理方式", "预期效果"], [
        ["Wi-Fi 配置无法保存", "增加 flash_test、errno 日志、写入后读回校验", "区分文件系统、FAL、Flash 硬件或旧固件问题"],
        ["连续点击按钮锁死", "呼噜启停加入生命周期互斥锁和线程退出等待", "避免重复 delete 线程造成 RT-Thread 断言"],
        ["语音和呼噜抢麦克风", "音频采集中心唯一打开 mic0，消费者独立缓冲", "小智语音和呼噜守护可同时运行"],
        ["报警声被识别为呼噜", "本地播放期间抑制呼噜告警，推理仍可继续更新状态", "减少设备自身声音导致的误报"],
        ["紧急状态传感器短暂掉线", "后端增加紧急状态宽限期", "SOS 播放期间前端不立即判定呼噜/环境离线"],
        ["旧模拟接口混淆", "正式路径统一迁移为 /hardware/*", "真实硬件数据优先，测试流程更清晰"],
    ], widths=[1.55, 2.65, 2.25])
    add_figure(doc, "image9.png", "图11 固件任务协同与处理流程示意", width=4.9)

    page_break(doc)
    add_heading(doc, "第四部分  算法与数据融合设计", 1)
    add_heading(doc, "4.1 雷达生命体征处理", 2)
    add_para(doc, "雷达端采集原始帧数据后，后端进行距离 FFT、静态杂波抑制、目标距离箱选择和相位提取。呼吸率通过相位信号频谱质量评估得到，心率结合信号分解和模型推理进行估计。为了减少运动或板子不稳定造成的误判，系统对低质量呼吸结果进行过滤，并在雷达离线、目标不在床或数据质量不足时不生成疑似呼吸暂停事件。")
    add_figure(doc, "image10.png", "图11 雷达距离 FFT 与目标提取处理效果", width=5.9)
    add_figure(doc, "image11.png", "图12 雷达相位与相位差分信号", width=5.9)
    add_figure(doc, "image12.png", "图13 雷达生命体征波形样例（一）", width=5.9)
    add_figure(doc, "image13.png", "图14 雷达生命体征波形样例（二）", width=5.9)
    add_heading(doc, "4.2 呼噜边缘检测模型", 2)
    add_para(doc, "呼噜检测模型部署在 M55 侧，模型输入为共享麦克风采集的单声道 PCM 片段。固件以 2 秒滑动窗口进行推理，输出 snore_score 和 detected 标志；本地 UI 显示推理结果，后端每秒接收心跳，用于前端趋势图和融合算法。模型采用 int8 数据表示，减少 RAM 与 Flash 占用，适合边缘设备长期运行。")
    add_figure(doc, "image19.png", "图15 呼噜/生命体征模型推理流程示意", width=4.6)
    add_heading(doc, "4.3 疑似睡眠呼吸暂停融合", 2)
    add_para(doc, "单独的雷达呼吸率异常或单独呼噜声都不直接判定为呼吸暂停。后端在 30-60 秒窗口内观察目标在床、雷达在线、呼吸率连续低或丢失，以及呼噜增强/恢复性声音证据，满足条件后生成 suspected_apnea 事件。该事件仅作为看护辅助提示，不计算医学意义上的 AHI。")
    add_figure(doc, "diagram_apnea_fusion.png", "图16 雷达 + 呼噜融合检测疑似呼吸暂停流程", width=6.1)
    add_table(doc, ["输入条件", "单独出现时处理", "与其他证据融合后"], [
        ["呼吸率连续偏低/丢失", "只标记候选，不立即报警", "若目标在床且伴随呼噜/恢复声，生成 suspected_apnea"],
        ["呼噜强度升高", "只生成呼噜扰动事件", "若前后出现低呼吸窗口，提高暂停置信度"],
        ["目标不在床", "不生成呼吸暂停", "仅提示离床/无人"],
        ["雷达离线或未静止", "不生成呼吸暂停", "提示设备状态，避免误判"],
        ["紧急语音/IMU", "直接 critical SOS", "不依赖呼吸暂停融合算法"],
    ], widths=[1.75, 2.25, 2.45])
    add_heading(doc, "4.4 睡眠环境评分", 2)
    add_para(doc, "环境分析页综合温度、湿度和环境声音强度生成睡眠环境评分。温度建议区间为 18-24℃，湿度建议区间为 40%-60%RH，声音强度以 dBFS 相对值表示。评分并非医学结论，而是帮助看护人员快速判断夜间环境是否可能影响入睡和睡眠稳定。")

    page_break(doc)
    add_heading(doc, "第五部分  后端、前端与 Android App", 1)
    add_heading(doc, "5.1 真实后端接口设计", 2)
    add_para(doc, "项目已以真实开发板数据为主，后端只保留正式硬件接口。M55 固件通过 /hardware/snore-heartbeat、/hardware/environment-heartbeat 和 /emergency 等接口上报，雷达板通过 UDP 发送生命体征数据。后端统一生成 /status、/timeline 和 /sleep/overview，供 Web 与 Android 读取。")
    add_figure(doc, "diagram_backend_interfaces.png", "图17 真实后端接口与核心数据对象", width=6.1)
    add_table(doc, ["接口", "来源/调用端", "作用"], [
        ["/hardware/snore-session/start", "M55 呼噜守护", "标记呼噜监测开始，前端显示呼噜守护在线"],
        ["/hardware/snore-heartbeat", "M55 呼噜检测任务", "上报 snore_score、dbfs 和 detected 状态"],
        ["/hardware/environment-heartbeat", "M55 环境任务", "上报温度、湿度、sensor_ok 和环境板在线状态"],
        ["/emergency", "M55 语音/IMU", "创建 critical 紧急事件"],
        ["/emergency/resolve", "开发板/Web/App", "解除 active emergency 并记录处理来源"],
        ["/sleep/overview", "Web/App", "聚合生命体征、事件流、设备状态和环境数据"],
    ], widths=[2.25, 1.7, 2.55])
    add_heading(doc, "5.2 Web 前端设计", 2)
    add_para(doc, "Web 前端以看护人员为主要使用对象，页面强调简洁、实时和可处理。睡眠驾驶舱展示实时生命体征、设备状态和事件流；看护预警中心按优先级给出处理动作；环境分析页展示温湿度、声音、评分和建议；历史数据页用于回看趋势。")
    add_figure(doc, "image27.png", "图18 生命体征与趋势展示界面示意", width=5.8)
    add_figure(doc, "image30.png", "图19 综合状态与看护页面示意", width=5.8)
    add_figure(doc, "image32.png", "图20 告警事件列表与处理界面示意", width=5.8)
    add_heading(doc, "5.3 Android 监护人 App", 2)
    add_para(doc, "Android App 不依赖 BLE 近场连接，而是绑定真实后端地址并通过前台服务持续轮询。这样监护人只要能访问后端，就可以在手机后台或锁屏状态下接收语音 SOS、IMU 摇晃和疑似呼吸暂停告警。App 内提供最近报警列表、详情和已处理按钮。")
    add_figure(doc, "image33.jpeg", "图21 Android App 报警列表界面示意", width=2.4)
    add_figure(doc, "image35.png", "图22 Android App 登录/绑定界面示意", width=2.4)

    page_break(doc)
    add_heading(doc, "第六部分  工程实现与测试结果", 1)
    add_heading(doc, "6.1 工程成果", 2)
    add_para(doc, "当前工程已形成由真实硬件、真实后端、Web 看护端和 Android App 组成的完整样机。M55 固件可编译并生成 rtthread.hex；前端可通过 npm run build 构建；后端关键 unittest 已通过；文档和测试清单已整理为真实数据优先。")
    add_figure(doc, "image20.jpeg", "图23 雷达与开发板实物连接示意（一）", width=5.4)
    add_figure(doc, "image22.jpeg", "图24 雷达与开发板实物连接示意（二）", width=5.4)
    add_figure(doc, "image23.jpeg", "图25 开发板、雷达与供电连接示意", width=5.4)
    add_figure(doc, "image24.jpeg", "图26 雷达模块安装与测试场景示意", width=5.4)
    add_heading(doc, "6.2 已完成功能清单", 2)
    add_table(doc, ["类别", "已完成功能", "验证方式"], [
        ["开发板", "小智语音、SOS 页面、解除按钮、报警声、闹钟界面", "串口日志、屏幕交互、真实后端事件"],
        ["呼噜守护", "共享麦克风、模型推理、暂停/恢复、后端心跳", "开发板 UI 与 /hardware/snore-heartbeat"],
        ["环境监测", "AHT20 温湿度、环境评分、趋势图", "/hardware/environment-heartbeat 与环境分析页"],
        ["雷达监测", "心率、呼吸率、距离、在床状态", "UDP 数据、驾驶舱生命体征卡片"],
        ["融合告警", "语音 SOS、IMU SOS、疑似呼吸暂停", "事件流、预警中心、Android 通知"],
        ["数据应用", "睡眠驾驶舱、预警中心、历史数据、Android App", "前端构建与实机联调"],
    ], widths=[1.2, 3.0, 2.2])
    add_heading(doc, "6.3 关键测试结果", 2)
    add_table(doc, ["序号", "测试项", "预期结果", "当前结果"], [
        ["1", "M55 固件编译", "可生成固件镜像", "scons 构建通过，生成 rtthread.hex"],
        ["2", "前端构建", "生产构建成功", "npm run build 通过"],
        ["3", "后端接口测试", "真实硬件接口可用", "关键 unittest 通过"],
        ["4", "SOS 联动", "开发板、后端、前端、App 同步", "语音和 IMU 共用 emergency 闭环"],
        ["5", "呼噜与语音并行", "不抢占麦克风", "共享音频中心分发 PCM"],
        ["6", "环境趋势", "温湿度可进入前端折线图", "环境分析页展示评分与建议"],
        ["7", "旧模拟路径清理", "真实数据为主", "/mock 相关正式链路已删除并迁移到 /hardware"],
        ["8", "文档与测试清单", "可复现实机测试", "已整理真实环境测试清单"],
    ], widths=[0.45, 1.55, 2.05, 2.3])
    add_heading(doc, "6.4 系统边界与安全说明", 2)
    add_callout(doc, "非医疗诊断声明", "本作品用于课程设计、比赛展示和看护辅助研究，不用于医疗诊断。疑似呼吸暂停、心率/呼吸异常等提示需要结合现场观察和专业设备复核。")
    add_para(doc, "在紧急事件处理上，系统 v1 不自动拨打电话或发送短信，而是通过开发板报警、Web 预警和 Android 通知提醒监护人确认。这一设计可以降低误报造成的外部扰动，同时保留完整的事件记录和人工处理入口。")

    add_heading(doc, "第七部分  总结与展望", 1)
    add_heading(doc, "7.1 项目总结", 2)
    add_para(doc, "本项目从单一雷达监测逐步扩展为多模态睡眠看护系统，工程难点集中在真实硬件链路稳定性、多任务资源冲突、告警状态同步和页面展示简洁性。通过共享麦克风、真实硬件接口迁移、紧急状态闭环、Android 后台通知和雷达 + 呼噜融合算法，系统已具备较完整的展示和测试价值。")
    add_heading(doc, "7.2 可扩展之处", 2)
    add_bullets(doc, [
        "加入血氧、压力或床垫传感器，提高睡眠呼吸暂停判断的可靠性。",
        "完善长期睡眠报告和趋势分析，形成周/月维度的看护摘要。",
        "部署公网 HTTPS 或家庭网关，使 Android App 在外网也能接收报警。",
        "进一步利用 PSOC Edge NPU/DSP 加速音频模型和更多边缘 AI 任务。",
        "增加设备配置页，让关键词、报警音量、闹钟和后端地址更加易用。",
    ])
    add_heading(doc, "7.3 心得体会", 2)
    add_para(doc, "开发过程中，我们对 RT-Thread 多任务、嵌入式音频采集、毫米波雷达信号处理、前后端实时状态同步和 Android 后台通知都有了更完整的实践认识。看护系统不能只追求算法输出，还需要考虑用户能否理解、监护人能否及时处理、误报是否可控、界面是否足够简洁。系统最终选择强告警与人工确认结合的方式，体现了辅助看护系统应有的谨慎边界。")

    add_heading(doc, "第八部分  参考文献", 1)
    refs = [
        "[1] Infineon Technologies AG. PSOC Edge E84 Consumer Datasheet.",
        "[2] Infineon Technologies AG. 2026 嵌入式系统设计大赛英飞凌赛道说明.",
        "[3] RT-Thread Studio. Edgi Talk SDK and RT-Thread BSP documentation.",
        "[4] 薛毅松. 基于毫米波雷达的非接触式健康监测系统研究[D]. 电子科技大学, 2022.",
        "[5] Zhao W, Tong G. A deep learning based heart rate estimation method for millimeter wave radar[J]. Measurement, 2025.",
        "[6] 王敬凯, 秦董洪, 白凤波, 等. 语音识别与大语言模型融合技术研究综述[J]. 计算机工程与应用, 2025.",
        "[7] AASM. Sleep-related breathing disorders and sleep apnea clinical background materials.",
        "[8] Infineon Community and AIROC Wi-Fi/Bluetooth technical resources.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(ref)
        font_run(r, size=9.5)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
